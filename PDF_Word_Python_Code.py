# -*- coding: utf-8 -*-
"""
Created on Fri Dec  4 13:05:38 2020

@author: heart
"""

# Python to process PDF 
# Batch process PDF to word

import pdfplumber
from docx import Document
import os
from multiprocessing import Process
#from exceptions import PendingDeprecationWarning

def convertPdf(fileName):
    with pdfplumber.open(fileName) as pdf:
        print("Processing{0},total {1} pages...".format(fileName,len(pdf.pages)))
        content = ''
        baseName = fileName.split('.')[0]
        wordName = baseName + '.docx'
        flag = True
        if os.path.exists(wordName):
            os.remove(wordName)
        for i in range(len(pdf.pages)):
            print("Processing <{0}> page {1} ...".format(baseName,i))
            page = pdf.pages[i]
            if page.extract_text() == None:
                print("{0} is image so it cannot be processed".format(fileName))
                flag = False
                break
            page_content = '\n'.join(page.extract_text().split('\n')[:-1])
            content = content + page_content
            if os.path.exists(wordName):
                doc = Document(wordName)
            else:
                doc = Document()
            doc.add_paragraph(content)
            doc.save(wordName)
            content = ''
            print("<{0}>Page {1} process is done!".format(baseName,i))
        if flag:
            print("{0} completed，saved as {1}!!!!".format(fileName,wordName))

if __name__ == '__main__':
    for file in os.listdir('.'):
        if os.path.isfile(file) and file.split('.')[1] == 'pdf':
            p = Process(target=convertPdf,args=(file,))
            p.start()

